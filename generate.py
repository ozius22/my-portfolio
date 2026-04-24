import json
from docx import Document
from docx.shared import Pt

def generate_docx_exam(json_file):
    try:
        with open(json_file, 'r') as f:
            data = json.load(f)
        
        exam_data = data.get("exam", {})
        questions = exam_data.get("questions", [])
        
        # Initialize Document
        doc = Document()
        
        # Exam Title and Header
        doc.add_heading(exam_data.get("title", "Summative Exam"), 0)
        doc.add_paragraph(f"Time Limit: {exam_data.get('timeLimitMinutes')} minutes")
        doc.add_paragraph(exam_data.get("description", ""))
        doc.add_paragraph("_" * 50)

        for i, q in enumerate(questions, 1):
            # Question Heading (Number and Points)
            p = doc.add_paragraph()
            run = p.add_run(f"Question {i} ({q.get('points')} pts)")
            run.bold = True
            
            # Question Text
            doc.add_paragraph(q.get("question"))

            q_type = q.get("type")

            # 1. Choices (Single/Multiple)
            if q_type in ["single_choice", "multiple_select"]:
                for choice in q.get("choices", []):
                    doc.add_paragraph(f"□ {choice}", style='List Bullet')

            # 2. Matching
            elif q_type == "matching":
                doc.add_paragraph("Match the items on the left with the correct description on the right:")
                left = q.get("leftItems", [])
                right = q.get("rightItems", [])
                for l, r in zip(left, right):
                    doc.add_paragraph(f"___ {l:<15} | {r}")

            # 3. Categorization
            elif q_type == "categorization":
                cats = ", ".join(q.get("categories", []))
                doc.add_paragraph(f"Categories: {cats}")
                doc.add_paragraph(f"Items: {', '.join(q.get('items', []))}")

            # 4. Code Snippets
            if "snippet" in q:
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(q["snippet"])
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)

            # 5. Case Study Sub-items
            if q_type == "case_study":
                for sub in q.get("subItems", []):
                    doc.add_paragraph(f"   - {sub.get('question')}")
                    for choice in sub.get("choices", []):
                        doc.add_paragraph(f"     □ {choice}")

            doc.add_paragraph("") # Spacing

        # Save the file
        doc.save("Generated_WebDev_Exam.docx")
        print("Success! 'Generated_WebDev_Exam.docx' has been created.")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    generate_docx_exam('exam.json')